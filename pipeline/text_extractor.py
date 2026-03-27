import os
import re
import logging
import asyncio
import time
from typing import List, Optional
from docx import Document
import pandas as pd
import pdfplumber  # оставлен для таблиц
import fitz  # PyMuPDF для быстрого текста
from bs4 import BeautifulSoup
from config import MAX_PDF_PAGES, MAX_DOCX_PARAGRAPHS, MAX_CONCURRENT_EXTRACT, PROCESS_DOC, MAX_DEPTH

try:
    from config import FILES_KEYWORDS
except ImportError:
    FILES_KEYWORDS = []

logger = logging.getLogger(__name__)

# Поддерживаемые расширения (без архивов)
SUPPORTED_EXTENSIONS = (
    ".pdf", ".docx", ".txt", ".json", ".xlsx", ".xls", ".html"
) + ((".doc",) if PROCESS_DOC else ())

# ------------------------------------------------------------
# Проверка валидности DOCX
# ------------------------------------------------------------
def is_valid_docx(file_path: str) -> bool:
    """Проверяет, является ли файл корректным ZIP-архивом с [Content_Types].xml."""
    import zipfile
    try:
        with zipfile.ZipFile(file_path, 'r') as zf:
            return '[Content_Types].xml' in zf.namelist()
    except zipfile.BadZipFile:
        return False

# ------------------------------------------------------------
# Очистка текста
# ------------------------------------------------------------
def clean_text(text, short_line_threshold=10, preserve_structure: bool = False):
    """
    Очищает текст от лишних пробелов, пустых строк и отступов.

    Если preserve_structure=True — старается не ломать исходную разметку
    (заголовки, таблицы, списки), не склеивает короткие строки.
    """
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        line = line.rstrip()
        norm = re.sub(r'[ \t]+', ' ', line)
        cleaned.append(norm)

    if preserve_structure:
        return "\n".join(cleaned)

    non_empty = [l.strip() for l in cleaned if l.strip()]

    merged = []
    for line in non_empty:
        if not merged:
            merged.append(line)
        else:
            if len(line) <= short_line_threshold:
                merged[-1] += " " + line
            else:
                merged.append(line)
    return '\n'.join(merged)

# ------------------------------------------------------------
# Преобразование таблиц в Markdown
# ------------------------------------------------------------
def dataframe_to_markdown(df: pd.DataFrame) -> str:
    """Преобразует pandas DataFrame в строку Markdown-таблицы."""
    if df.empty:
        return ""
    header = "| " + " | ".join(str(col) for col in df.columns) + " |"
    separator = "|" + "|".join([" --- "] * len(df.columns)) + "|"
    rows = []
    for _, row in df.iterrows():
        row_str = [str(val) if pd.notna(val) else "" for val in row]
        rows.append("| " + " | ".join(row_str) + " |")
    return "\n".join([header, separator] + rows)

def table_to_markdown(table_data) -> str:
    """
    Универсальная функция: принимает таблицу в виде списка списков или DataFrame,
    возвращает Markdown-строку.
    """
    if isinstance(table_data, pd.DataFrame):
        return dataframe_to_markdown(table_data)
    elif isinstance(table_data, list) and all(isinstance(row, list) for row in table_data):
        if not table_data:
            return ""
        if len(table_data) > 1:
            df = pd.DataFrame(table_data[1:], columns=table_data[0])
        else:
            df = pd.DataFrame(table_data)
        return dataframe_to_markdown(df)
    else:
        return ""

# ------------------------------------------------------------
# Извлечение таблиц из PDF (оставлен pdfplumber, опционально)
# ------------------------------------------------------------
def extract_tables_from_pdf(pdf_path: str, max_pages: Optional[int] = None) -> List[str]:
    """Извлекает таблицы из PDF с помощью pdfplumber и возвращает список Markdown-строк."""
    tables_md = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                if max_pages and page_num >= max_pages:
                    break
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:
                        md = table_to_markdown(table)
                        if md:
                            tables_md.append(f"\n## Таблица со страницы {page_num+1}\n\n{md}\n")
    except Exception as e:
        logger.debug(f"Ошибка при извлечении таблиц из PDF {pdf_path}: {e}")
    return tables_md

def extract_tables_from_docx(docx_path: str) -> List[str]:
    """Извлекает таблицы из DOCX с помощью python-docx."""
    tables_md = []
    try:
        doc = Document(docx_path)
        for table_idx, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                data.append([cell.text.strip() for cell in row.cells])
            if data:
                md = table_to_markdown(data)
                if md:
                    tables_md.append(f"\n## Таблица {table_idx+1}\n\n{md}\n")
    except Exception as e:
        logger.debug(f"Ошибка при извлечении таблиц из DOCX {docx_path}: {e}")
    return tables_md

def extract_tables_from_excel(excel_path: str) -> List[str]:
    """Извлекает все листы Excel как таблицы в Markdown."""
    tables_md = []
    try:
        xl = pd.ExcelFile(excel_path)
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            if not df.empty:
                md = dataframe_to_markdown(df)
                if md:
                    tables_md.append(f"\n## Лист: {sheet_name}\n\n{md}\n")
    except Exception as e:
        logger.debug(f"Ошибка при извлечении таблиц из Excel {excel_path}: {e}")
    return tables_md

def extract_text_from_html(html_path: str) -> str:
    """Извлекает чистый текст из HTML-файла с помощью BeautifulSoup."""
    try:
        with open(html_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'lxml')
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text(separator='\n')
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return '\n'.join(lines)
    except Exception as e:
        logger.debug(f"Ошибка при извлечении текста из HTML {html_path}: {e}")
        return ""

# ------------------------------------------------------------
# Основная функция извлечения текста (для одного файла)
# ------------------------------------------------------------
def extract_text(file_path: str, max_pages: Optional[int] = None, extract_tables: bool = True) -> str:
    """
    Извлекает текст и таблицы (в Markdown) из файла.
    Возвращает содержимое, готовое для сохранения в .md.

    Параметры:
        extract_tables: извлекать ли таблицы из PDF (медленно, по умолчанию False)
    """
    if not os.path.isfile(file_path):
        logger.error(f"Файл не найден: {file_path}")
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    # Для JSON возвращаем блок кода
    if ext == '.json':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return f"```json\n{content}\n```"
        except Exception as e:
            logger.exception(f"Ошибка при чтении JSON {file_path}: {e}")
            raise

    # Для TXT просто читаем (может быть Markdown или обычный текст)
    if ext == '.txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.exception(f"Ошибка при чтении TXT {file_path}: {e}")
            raise

    # HTML
    if ext in ('.html', '.htm'):
        return extract_text_from_html(file_path)

    # PDF – теперь с PyMuPDF (быстро)
    if ext == '.pdf':
        pages_limit = max_pages if max_pages is not None else MAX_PDF_PAGES
        full_text = ""

        # 1. Извлекаем текст через PyMuPDF
        try:
            doc = fitz.open(file_path)
            # Если pages_limit задан, ограничиваем число страниц
            total_pages = len(doc)
            pages_to_read = min(total_pages, pages_limit) if pages_limit else total_pages
            for page_num in range(pages_to_read):
                page = doc[page_num]
                text = page.get_text("text")  # быстрый метод, без layout
                if text.strip():
                    full_text += f"\n\n## Страница {page_num + 1}\n\n"
                    full_text += text + "\n"
            doc.close()
        except Exception as e:
            logger.exception(f"Ошибка при чтении текста из PDF {file_path}: {e}")

        # 2. Извлекаем таблицы, если нужно (опционально, медленно)
        if extract_tables:
            tables_md = extract_tables_from_pdf(file_path, max_pages=pages_limit)
            if tables_md:
                full_text += "\n\n## Извлечённые таблицы\n\n"
                full_text += "\n".join(tables_md)

        return clean_text(full_text, preserve_structure=True)

    # DOCX
    if ext == '.docx':
        full_text_lines = []

        # Проверка валидности DOCX
        if not is_valid_docx(file_path):
            logger.warning(f"Файл {file_path} имеет расширение .docx, но не является валидным DOCX. "
                           f"Возможно, это старый .doc. Попытка извлечения через Word (если включено).")
            if PROCESS_DOC:
                try:
                    text = _extract_doc_via_word(file_path)
                    return clean_text(text, preserve_structure=True)
                except Exception as e:
                    logger.warning(f"Не удалось извлечь текст через Word: {e}")
                    return ""
            else:
                logger.warning("PROCESS_DOC отключён, пропускаем файл.")
                return ""

        # Основное извлечение из DOCX
        try:
            doc = Document(file_path)
            limit = max_pages if max_pages is not None else MAX_DOCX_PARAGRAPHS
            for p in doc.paragraphs[:limit]:
                text = p.text.strip()
                if not text:
                    continue
                style_name = (p.style.name or "").lower() if p.style else ""

                # Преобразуем стили Heading в Markdown-заголовки
                level = None
                if "heading" in style_name:
                    m = re.search(r'heading\s*(\d+)', style_name)
                    if m:
                        try:
                            level = int(m.group(1))
                        except ValueError:
                            level = None
                if level is not None and 1 <= level <= 6:
                    prefix = "#" * level
                    full_text_lines.append(f"{prefix} {text}")
                else:
                    full_text_lines.append(text)
        except Exception as e:
            logger.warning(f"Ошибка при чтении текста из DOCX {file_path}: {e}")
            full_text_lines = []

        full_text = "\n".join(full_text_lines) + "\n"

        # Таблицы
        tables_md = extract_tables_from_docx(file_path)
        if tables_md:
            full_text += "\n\n## Извлечённые таблицы\n\n"
            full_text += "\n".join(tables_md)

        return clean_text(full_text, preserve_structure=True)

    # Excel (.xlsx, .xls)
    if ext in ('.xlsx', '.xls'):
        tables_md = extract_tables_from_excel(file_path)
        if tables_md:
            full_text = "\n\n".join(tables_md)
            return full_text.strip()
        else:
            return ""

    # DOC (если включён)
    if ext == '.doc' and PROCESS_DOC:
        try:
            text = _extract_doc_via_word(file_path)
            return clean_text(text, preserve_structure=True)
        except Exception as e:
            logger.exception(f"Ошибка при чтении DOC {file_path}: {e}")
            raise

    raise ValueError(f"Неподдерживаемый формат файла: {ext}")

def _extract_doc_via_word(file_path):
    """Извлекает текст из .doc через Microsoft Word (COM automation)."""
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        raise RuntimeError(
            "Для обработки .doc нужен MS Word и пакет pywin32. "
            "Установите зависимость: pip install pywin32"
        )

    abs_path = os.path.abspath(file_path)
    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        try:
            return doc.Content.Text or ""
        finally:
            doc.Close(False)
            word.Quit()
    finally:
        pythoncom.CoUninitialize()

# ------------------------------------------------------------
# Асинхронная обёртка
# ------------------------------------------------------------
async def extract_text_async(file_path, max_pages=None, extract_tables=False):
    return await asyncio.to_thread(extract_text, file_path, max_pages, extract_tables)

# ------------------------------------------------------------
# Основная функция: конвертация папки с сохранением структуры
# ------------------------------------------------------------
async def process_folder(
    input_folder: str,
    output_folder: str,
    max_pages: Optional[int] = None,
    concurrency: Optional[int] = None,
    max_depth: Optional[int] = MAX_DEPTH,
    keywords: Optional[List[str]] = FILES_KEYWORDS,
    extract_tables: bool = False
) -> List[str]:
    """
    Обходит папку input_folder, извлекает текст из всех поддерживаемых файлов
    и сохраняет результат в output_folder с расширением .md, сохраняя структуру каталогов.
    В начало каждого файла добавляется заголовок с именем исходного файла.

    Аргументы:
        input_folder (str): путь к исходной папке (уже разархивированной)
        output_folder (str): путь к папке, куда будут сохранены .md файлы
        max_pages (int, optional): максимальное количество страниц для PDF
        concurrency (int, optional): количество параллельных задач (по умолчанию из конфига)
        max_depth (int, optional): максимальная глубина вложенности (0 - только корень)
        keywords (List[str], optional): список ключевых слов для фильтрации по имени файла.
                                        Если список пуст или None, фильтрация отключена.
        extract_tables (bool): извлекать ли таблицы из PDF (медленно, по умолчанию False)
    """
    input_folder = os.path.abspath(input_folder)
    output_folder = os.path.abspath(output_folder)

    if concurrency is None:
        concurrency = MAX_CONCURRENT_EXTRACT

    # Собираем список всех файлов, которые будем обрабатывать
    all_files = []
    for root, _, files in os.walk(input_folder):
        # Вычисляем глубину root относительно input_folder
        rel_root = os.path.relpath(root, input_folder)
        if rel_root == '.':
            depth = 0
        else:
            depth = rel_root.count(os.sep) + 1

        if max_depth is not None and depth > max_depth:
            continue   # пропускаем каталоги глубже max_depth

        for name in files:
            _, ext = os.path.splitext(name)
            if ext.lower() in SUPPORTED_EXTENSIONS:
                all_files.append(os.path.join(root, name))

    total_files = len(all_files)
    logger.info("Найдено поддерживаемых файлов: %s", total_files)

    sem = asyncio.Semaphore(concurrency)
    created_files = []
    processed_count = 0
    lock = asyncio.Lock()

    async def _handle_file(src_path: str, out_path: str, file_idx: int):
        nonlocal processed_count
        start_time = time.time()
        logger.info("[%d/%d] Начало обработки: %s", file_idx + 1, total_files, src_path)
        async with sem:
            try:
                # Фильтрация по ключевым словам в имени файла
                if keywords:
                    filename = os.path.basename(src_path)
                    name_without_ext = os.path.splitext(filename)[0]
                    name_lower = name_without_ext.lower()
                    if not any(kw.lower() in name_lower for kw in keywords):
                        logger.info("[%d/%d] Имя файла не содержит ключевых слов, пропущен: %s",
                                    file_idx + 1, total_files, src_path)
                        return

                # Извлечение текста (только если прошёл фильтр)
                text = await asyncio.to_thread(extract_text, src_path, max_pages, extract_tables)

                # Сохраняем результат
                filename = os.path.basename(src_path)
                header = f"# Файл: {filename}\n\n"
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(header + text)

                async with lock:
                    created_files.append(out_path)
                    processed_count += 1
                elapsed = time.time() - start_time
                logger.info("[%d/%d] Успешно обработан: %s (время: %.2f сек)",
                            processed_count, total_files, src_path, elapsed)
            except Exception as e:
                elapsed = time.time() - start_time
                logger.exception("[%d/%d] Ошибка обработки %s (время: %.2f сек): %s",
                                 file_idx + 1, total_files, src_path, elapsed, e)

    tasks = []
    for idx, src_path in enumerate(all_files):
        name = os.path.basename(src_path)
        root = os.path.dirname(src_path)
        rel = os.path.relpath(root, input_folder)
        out_dir = output_folder if rel == "." else os.path.join(output_folder, rel)
        os.makedirs(out_dir, exist_ok=True)

        base = os.path.splitext(name)[0]
        out_path = os.path.join(out_dir, base + ".md")
        tasks.append(asyncio.create_task(_handle_file(src_path, out_path, idx)))

    if tasks:
        await asyncio.gather(*tasks)

    logger.info("Обработано файлов: %s, сохранено в %s", len(created_files), output_folder)
    return created_files

# ------------------------------------------------------------
# Для тестирования
# ------------------------------------------------------------
def main():
    logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s: %(message)s")
    # Пример использования:
    asyncio.run(process_folder(
        input_folder=r"F:\Python projects\Aut\storage\tenders\downloaded_files\90918705",
        output_folder="output_md",
        max_pages=10,           # ограничиваем страницы для ускорения
        concurrency=20,         # увеличиваем параллелизм
        max_depth=3,
        keywords=FILES_KEYWORDS,
        extract_tables=True
    ))

if __name__ == "__main__":
    main()
